import os
from docx import Document

def clear_header_footer_safe():
    # ================= 配置区 =================
    # 输出文件夹的名字
    OUTPUT_FOLDER_NAME = "无页眉页脚结果"
    # ========================================

    current_folder = os.getcwd()
    output_path = os.path.join(current_folder, OUTPUT_FOLDER_NAME)

    # 1. 创建输出文件夹 (如果不存在)
    if not os.path.exists(output_path):
        os.makedirs(output_path)
        print(f"📂 已新建文件夹: {OUTPUT_FOLDER_NAME}")

    # 2. 扫描文件 (排除临时文件)
    files = [f for f in os.listdir(current_folder) if f.endswith(".docx") and not f.startswith("~$")]

    if not files:
        print("⚠️  当前目录下没有找到 .docx 文件。")
        return

    print(f"🚀 扫描到 {len(files)} 个文件，开始清理页眉页脚...\n")

    success_count = 0

    for filename in files:
        input_file = os.path.join(current_folder, filename)
        output_file = os.path.join(output_path, filename)

        try:
            print(f"处理中: {filename} ... ", end="")
            
            # 打开文档
            doc = Document(input_file)

            # 遍历每一个“节” (Section)
            for section in doc.sections:
                # 3. 清理页眉 (Header) - 包含普通页、首页、偶数页
                headers = [section.header, section.first_page_header, section.even_page_header]
                for h in headers:
                    if h: # 确保对象存在
                        for p in h.paragraphs:
                            p.text = "" # 清空文字

                # 4. 清理页脚 (Footer) - 包含普通页、首页、偶数页
                footers = [section.footer, section.first_page_footer, section.even_page_footer]
                for f in footers:
                    if f: # 确保对象存在
                        for p in f.paragraphs:
                            p.text = "" # 清空文字

            # 5. 保存到新文件夹
            doc.save(output_file)
            print("✅ 完成")
            success_count += 1

        except Exception as e:
            print(f"❌ 失败! 原因: {e}")

    print("\n" + "="*30)
    print(f"🎉 全部处理完毕！成功: {success_count} / {len(files)}")
    print(f"📁 干净的文档已保存在: {OUTPUT_FOLDER_NAME}")

if __name__ == "__main__":
    clear_header_footer_safe()